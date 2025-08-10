from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from typing import List, Optional
import os

class DocumentProcessor:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def process_text(self, text: str, metadata: Optional[dict] = None) -> List[Document]:
        """
        Process raw text into chunks for vector storage
        """
        if metadata is None:
            metadata = {}
        
        # Create a document
        doc = Document(page_content=text, metadata=metadata)
        
        # Split into chunks
        chunks = self.text_splitter.split_documents([doc])
        
        return chunks
    
    def process_file(self, file_path: str, metadata: Optional[dict] = None) -> List[Document]:
        """
        Process a file (.txt, .md, .pdf, .docx) into chunks
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()

        if ext in (".txt", ".md"):
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        elif ext == ".pdf":
            try:
                from pypdf import PdfReader
            except ImportError as e:
                raise ImportError("pypdf is required to process PDF files. Please install it.") from e
            reader = PdfReader(file_path)
            pages_text = []
            for page in reader.pages:
                try:
                    pages_text.append(page.extract_text() or "")
                except Exception:
                    pages_text.append("")
            text = "\n\n".join(pages_text)
        elif ext == ".docx":
            try:
                from docx import Document as DocxDocument
            except ImportError as e:
                raise ImportError("python-docx is required to process DOCX files. Please install it.") from e
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text]
            text = "\n".join(paragraphs)
        else:
            raise ValueError(f"Unsupported file extension: {ext}")

        if metadata is None:
            metadata = {}

        metadata['source'] = file_path
        metadata['file_name'] = os.path.basename(file_path)

        return self.process_text(text, metadata)